import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_and_convert(md_text, output_file="output.docx", is_proof2=False):
    doc = Document()
    
    # 1. 状态机：将文本按照 Block 划分
    blocks = []
    current_block = {"type": None, "lines": []}
    
    for line in md_text.split('\n'):
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # 【修改点1】：探测区块起始点，改为正则匹配以 1 到多个 # 开头的行
        if re.match(r'^#+', line_stripped):
            if current_block["type"]:
                blocks.append(current_block)
            current_block = {"type": "heading", "lines": [line_stripped]}
        elif line_stripped.startswith('**['):
            if current_block["type"]:
                blocks.append(current_block)
            current_block = {"type": "regular", "lines": [line_stripped]}
        else:
            if current_block["type"]:
                current_block["lines"].append(line_stripped)
                
    if current_block["type"]:
        blocks.append(current_block)
        
    # 2. 根据区块类型分配处理逻辑
    for block in blocks:
        if block["type"] == "heading":
            process_heading_block(block["lines"], doc, is_proof2)
        elif block["type"] == "regular":
            process_regular_block(block["lines"], doc, is_proof2)
            
    doc.save(output_file)
    print(f"转换成功！文件已保存至：{output_file}")

def process_heading_block(lines, doc, is_proof2=False):
    """处理标题区块"""
    heading_text, en_text, block_id, orig_trans, proof, comments = "", "", "", "", "", ""
    heading_level = 1 # 默认设为一级标题
    
    # 提取字段
    for line in lines:
        # 【修改点2】：动态解析 # 的数量和标题正文
        if re.match(r'^#+', line):
            # 捕获组1：所有的#号；捕获组2：#号后面的所有文本（允许中间有空格或没空格）
            match = re.match(r'^(#{1,6})\s*(.*)', line)
            heading_level = len(match.group(1))
            # Word的默认标题级别最大为9，我们这里做一个安全限制
            heading_level = min(heading_level, 9) 
            heading_text = match.group(2).strip()
            
        elif line.startswith('*') and '`[' in line:
            m_en = re.search(r'\*(.*?)\*', line)
            if m_en: en_text = m_en.group(1)
            m_id = re.search(r'`\[(.*?)\]`', line)
            if m_id: block_id = m_id.group(1)
            
        elif line.startswith('> 原始译文:'):
            # 将原始译文中的 # 也顺便清理掉，保持正文清爽
            orig_trans = re.sub(r'^> 原始译文:\s*#*\s*', '', line).strip()
            
        elif line.startswith('> 校对:'):
            proof = line.replace('> 校对:', '').strip()
            if proof.startswith('**') and proof.endswith('**'):
                proof = proof[2:-2].strip()
                
        elif line.startswith('> 标题建议:') or line.startswith('> 建议:') or line.startswith('> *建议:'):
            comments = line.replace('> 标题建议:', '').replace('> 建议:', '').replace('> *建议:', '').strip()
            if comments.endswith('*'): comments = comments[:-1].strip()
            
    # 【需求2】按照新顺序重组标题区块内容
    # 1. 标题 -> 做真实标题，并动态传入 level 级别
    if heading_text:
        doc.add_heading(heading_text, level=heading_level)
    # 2. ID -> 正文格式
    if block_id:
        doc.add_paragraph(f"ID: [{block_id}]")
    # 3. 原文 -> 正文格式
    if en_text:
        doc.add_paragraph(f"原文: {en_text}")
    # 4. 原始译文 -> 正文格式
    if orig_trans:
        doc.add_paragraph(f"原始译文: {orig_trans}")
    # 5. 校对 -> 正文格式
    proof_label = "二校" if is_proof2 else "一校"
    if proof:
        doc.add_paragraph(f"{proof_label}: {proof}")
    # 6. 建议 -> 引用格式
    if comments:
        comment_label = "二校建议" if is_proof2 else "一校建议"
        p = doc.add_paragraph(f"{comment_label}: {comments}")
        p.style = 'Intense Quote'
        
    # 【需求1】每个 block 之间增加两个换行
    doc.add_paragraph()
    doc.add_paragraph()

def process_regular_block(lines, doc, is_proof2=False):
    """处理正文与表格区块"""
    block_id = ""
    content = {}
    current_key = None
    
    for line in lines:
        if line.startswith('**[') and line.endswith(']**'):
            block_id = line[2:-2]
        elif line.startswith('> 原文:'):
            current_key = '原文'
            content[current_key] = line[len('> 原文:'):].strip()
        elif line.startswith('> 原始译文:'):
            current_key = '原始译文'
            content[current_key] = line[len('> 原始译文:'):].strip()
        elif line.startswith('> 校对:'):
            current_key = '校对'
            text = line[len('> 校对:'):].strip()
            if text.startswith('**') and text.endswith('**'):
                text = text[2:-2].strip()
            content[current_key] = text
        elif line.startswith('> *建议:') or line.startswith('> 建议:'):
            current_key = '注释'
            text = line.replace('> *建议:', '').replace('> 建议:', '').strip()
            if text.endswith('*'): text = text[:-1].strip()
            content[current_key] = text
        else:
            if current_key:
                content[current_key] += " " + line
                
    if block_id:
        doc.add_paragraph(f"ID: {block_id}")
        
    # 处理原文（如果没有原文则原文为一校）
    if '原文' in content:
        text_val = content['原文']
        if '<table' in text_val.lower():
            doc.add_paragraph("原文:")
            handle_html_table(text_val, doc)
        else:
            doc.add_paragraph(f"原文: {text_val}")
    elif is_proof2 and '校对' in content:
        # 二校时如果没有原文，原文为一校
        text_val = content['校对']
        if '<table' in text_val.lower():
            doc.add_paragraph("原文:")
            handle_html_table(text_val, doc)
        else:
            doc.add_paragraph(f"原文: {text_val}")
    
    # 处理原始译文
    if '原始译文' in content:
        text_val = content['原始译文']
        if '<table' in text_val.lower():
            doc.add_paragraph("原始译文:")
            handle_html_table(text_val, doc)
        else:
            doc.add_paragraph(f"原始译文: {text_val}")
    
    # 处理校对
    proof_label = "二校" if is_proof2 else "一校"
    if '校对' in content:
        text_val = content['校对']
        if '<table' in text_val.lower():
            doc.add_paragraph(f"{proof_label}:")
            handle_html_table(text_val, doc)
        else:
            doc.add_paragraph(f"{proof_label}: {text_val}")
    
    # 处理建议
    if '注释' in content:
        comment_label = "二校建议" if is_proof2 else "一校建议"
        p = doc.add_paragraph(f"{comment_label}: {content['注释']}")
        p.style = 'Intense Quote'
        
    doc.add_paragraph()
    doc.add_paragraph()

def handle_html_table(html_str, doc):
    """解析并绘制原生 Word 表格"""
    soup = BeautifulSoup(html_str, 'html.parser')
    tables = soup.find_all('table')
    
    for table in tables:
        rows = table.find_all('tr')
        if not rows: continue
        
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = 'Table Grid'
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    doc_cell = word_table.cell(row_idx, col_idx)
                    doc_cell.text = cell.get_text(strip=True)
                    for paragraph in doc_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

